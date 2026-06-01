"""Tests for agentnexus.utils.docx.analyzer — DOCX 约束分析器。"""

from __future__ import annotations

import os
import tempfile
import zipfile
from pathlib import Path

import pytest


# ── 辅助：创建测试 DOCX 文件 ──

def _create_simple_docx(path: str, paragraphs: list[str] | None = None) -> None:
    """创建一个简单的 DOCX 文件用于测试。"""
    try:
        from docx import Document
        from docx.shared import Pt, Mm
    except ImportError:
        pytest.skip("python-docx 未安装")

    doc = Document()

    # 设置页面边距
    section = doc.sections[0]
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(30)
    section.right_margin = Mm(30)

    if paragraphs is None:
        paragraphs = ["测试段落一", "测试段落二", "测试段落三"]

    for text in paragraphs:
        doc.add_paragraph(text)

    doc.save(path)


def _create_docx_with_table(path: str) -> None:
    """创建包含表格的 DOCX 文件。"""
    try:
        from docx import Document
        from docx.shared import Mm
    except ImportError:
        pytest.skip("python-docx 未安装")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(30)
    section.right_margin = Mm(30)

    doc.add_paragraph("表格标题")

    table = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            table.rows[r].cells[c].text = f"R{r}C{c}"

    doc.add_paragraph("表格后的段落")
    doc.save(path)


def _create_docx_with_styles(path: str) -> None:
    """创建包含多种样式的 DOCX 文件。"""
    try:
        from docx import Document
        from docx.shared import Pt, Mm
    except ImportError:
        pytest.skip("python-docx 未安装")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(30)
    section.right_margin = Mm(30)

    doc.add_heading("一级标题", level=1)
    doc.add_paragraph("正文段落内容")
    doc.add_heading("二级标题", level=2)
    doc.add_paragraph("另一段正文")
    doc.add_heading("三级标题", level=3)
    doc.add_paragraph("最后一段")

    doc.save(path)


# ── 测试：基础分析功能 ──

class TestAnalyzeBasic:
    """基础分析功能测试。"""

    def test_analyze_simple_docx(self, tmp_path):
        """测试分析简单 DOCX 文件。"""
        from agentnexus.utils.docx.analyzer import analyze

        docx_path = str(tmp_path / "test.docx")
        _create_simple_docx(docx_path)

        result = analyze(docx_path)

        assert result.file_path == str(Path(docx_path).absolute())
        assert result.file_size_bytes > 0
        assert len(result.sections) >= 1
        assert result.default_page.width_mm > 0
        assert result.default_page.height_mm > 0

    def test_analyze_nonexistent_file(self):
        """测试分析不存在的文件。"""
        from agentnexus.utils.docx.analyzer import analyze

        with pytest.raises(FileNotFoundError):
            analyze("/nonexistent/file.docx")

    def test_analyze_invalid_file(self, tmp_path):
        """测试分析无效文件。"""
        from agentnexus.utils.docx.analyzer import analyze

        bad_path = str(tmp_path / "bad.docx")
        Path(bad_path).write_text("not a docx file")

        with pytest.raises(zipfile.BadZipFile):
            analyze(bad_path)


# ── 测试：页面约束 ──

class TestPageConstraints:
    """页面约束分析测试。"""

    def test_default_page_dimensions(self, tmp_path):
        """测试默认页面尺寸。"""
        from agentnexus.utils.docx.analyzer import analyze

        docx_path = str(tmp_path / "test.docx")
        _create_simple_docx(docx_path)

        result = analyze(docx_path)
        page = result.default_page

        # 默认纸张尺寸（python-docx 默认 Letter: 215.9mm × 279.4mm）
        assert 200 < page.width_mm < 220
        assert 270 < page.height_mm < 300

    def test_custom_margins(self, tmp_path):
        """测试自定义边距。"""
        from agentnexus.utils.docx.analyzer import analyze

        docx_path = str(tmp_path / "test.docx")
        _create_simple_docx(docx_path)

        result = analyze(docx_path)
        page = result.default_page

        # 我们设置的边距是 25mm 上下，30mm 左右
        assert 24 < page.margin_top_mm < 26
        assert 24 < page.margin_bottom_mm < 26
        assert 29 < page.margin_left_mm < 31
        assert 29 < page.margin_right_mm < 31

    def test_usable_width(self, tmp_path):
        """测试可用宽度计算。"""
        from agentnexus.utils.docx.analyzer import analyze

        docx_path = str(tmp_path / "test.docx")
        _create_simple_docx(docx_path)

        result = analyze(docx_path)
        page = result.default_page

        # 可用宽度 = 页面宽度 - 左右边距
        expected = page.width_mm - page.margin_left_mm - page.margin_right_mm
        assert abs(page.usable_width_mm - expected) < 0.1
        assert page.usable_width_chars > 0


# ── 测试：表格分析 ──

class TestTableAnalysis:
    """表格约束分析测试。"""

    def test_analyze_table_structure(self, tmp_path):
        """测试分析表格结构。"""
        from agentnexus.utils.docx.analyzer import analyze

        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_table(docx_path)

        result = analyze(docx_path)

        assert len(result.tables) >= 1
        tbl = result.tables[0]
        assert tbl.row_count == 3
        assert tbl.col_count == 3

    def test_table_column_widths(self, tmp_path):
        """测试表格列宽分析。"""
        from agentnexus.utils.docx.analyzer import analyze

        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_table(docx_path)

        result = analyze(docx_path)
        tbl = result.tables[0]

        assert len(tbl.columns) == 3
        for col in tbl.columns:
            assert col.width_mm >= 0


# ── 测试：样式分析 ──

class TestStyleAnalysis:
    """样式约束分析测试。"""

    def test_analyze_styles(self, tmp_path):
        """测试分析样式定义。"""
        from agentnexus.utils.docx.analyzer import analyze

        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_styles(docx_path)

        result = analyze(docx_path)

        # 应该有样式定义
        assert len(result.styles) > 0

    def test_heading_styles_detected(self, tmp_path):
        """测试标题样式被正确识别。"""
        from agentnexus.utils.docx.analyzer import analyze

        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_styles(docx_path)

        result = analyze(docx_path)

        # 检查是否有标题样式
        headings = [s for s in result.styles.values() if s.is_heading]
        assert len(headings) > 0


# ── 测试：摘要输出 ──

class TestSummary:
    """约束摘要输出测试。"""

    def test_summary_contains_page_info(self, tmp_path):
        """测试摘要包含页面信息。"""
        from agentnexus.utils.docx.analyzer import analyze_to_string

        docx_path = str(tmp_path / "test.docx")
        _create_simple_docx(docx_path)

        summary = analyze_to_string(docx_path)

        assert "页面设置" in summary
        assert "mm" in summary
        assert "可用宽度" in summary

    def test_summary_contains_rules(self, tmp_path):
        """测试摘要包含约束规则。"""
        from agentnexus.utils.docx.analyzer import analyze_to_string

        docx_path = str(tmp_path / "test.docx")
        _create_simple_docx(docx_path)

        summary = analyze_to_string(docx_path)

        assert "关键规则" in summary
        assert "表格" in summary
        assert "图片" in summary

    def test_analyze_to_string_error_handling(self):
        """测试错误处理。"""
        from agentnexus.utils.docx.analyzer import analyze_to_string

        result = analyze_to_string("/nonexistent/file.docx")
        assert "错误" in result


# ── 测试：段落读取 ──

class TestReadParagraphs:
    """段落元数据读取测试。"""

    def test_read_paragraphs_metadata(self, tmp_path):
        """测试读取段落元数据。"""
        from agentnexus.utils.docx.analyzer import read_paragraphs

        docx_path = str(tmp_path / "test.docx")
        _create_simple_docx(docx_path, ["段落一", "段落二", "段落三"])

        paragraphs = read_paragraphs(docx_path)

        assert len(paragraphs) >= 3
        for p in paragraphs:
            assert "index" in p
            assert "text" in p
            assert "style" in p
