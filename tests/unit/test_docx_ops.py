"""Tests for agentnexus.utils.docx.ops — DOCX 文档操作函数。"""

from __future__ import annotations

import os
from pathlib import Path

import pytest


# ── 辅助：创建测试 DOCX 文件 ──

def _create_docx(path: str, paragraphs: list[str] | None = None) -> None:
    """创建测试 DOCX 文件。"""
    try:
        from docx import Document
        from docx.shared import Mm
    except ImportError:
        pytest.skip("python-docx 未安装")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(30)
    section.right_margin = Mm(30)

    if paragraphs is None:
        paragraphs = ["第一段内容", "第二段内容", "第三段内容"]

    for text in paragraphs:
        doc.add_paragraph(text)

    doc.save(path)


def _create_docx_with_table(path: str) -> None:
    """创建包含表格的 DOCX 文件。"""
    try:
        from docx import Document
        from docx.shared import Mm
    except ImportError:
        pytest.skip("python-docx 未安装")

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Mm(25)
    section.bottom_margin = Mm(25)
    section.left_margin = Mm(30)
    section.right_margin = Mm(30)

    doc.add_paragraph("标题段落")

    table = doc.add_table(rows=3, cols=3)
    for r in range(3):
        for c in range(3):
            table.rows[r].cells[c].text = f"R{r}C{c}"

    doc.add_paragraph("尾部段落")
    doc.save(path)


# ── 测试：读取操作 ──

class TestReadOperations:
    """读取操作测试。"""

    def test_read_paragraphs(self, tmp_path):
        """测试读取段落。"""
        from agentnexus.utils.docx.ops import read_paragraphs

        docx_path = str(tmp_path / "test.docx")
        _create_docx(docx_path, ["段落A", "段落B", "段落C"])

        result = read_paragraphs(docx_path)

        assert len(result) >= 3
        texts = [p["text"] for p in result]
        assert "段落A" in texts
        assert "段落B" in texts
        assert "段落C" in texts

    def test_read_paragraphs_format(self, tmp_path):
        """测试读取段落格式信息。"""
        from agentnexus.utils.docx.ops import read_paragraphs

        docx_path = str(tmp_path / "test.docx")
        _create_docx(docx_path)

        result = read_paragraphs(docx_path)

        for p in result:
            assert "index" in p
            assert "text" in p
            assert "style" in p
            assert "font_name" in p
            assert "font_size_pt" in p
            assert "bold" in p

    def test_read_tables(self, tmp_path):
        """测试读取表格。"""
        from agentnexus.utils.docx.ops import read_tables

        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_table(docx_path)

        result = read_tables(docx_path)

        assert len(result) >= 1
        tbl = result[0]
        assert tbl["rows"] == 3
        assert tbl["cols"] == 3
        assert len(tbl["cells"]) == 3
        assert tbl["cells"][0][0] == "R0C0"

    def test_read_tables_empty(self, tmp_path):
        """测试读取无表格的文档。"""
        from agentnexus.utils.docx.ops import read_tables

        docx_path = str(tmp_path / "test.docx")
        _create_docx(docx_path)

        result = read_tables(docx_path)
        assert result == []


# ── 测试：文本替换 ──

class TestReplaceText:
    """文本替换测试。"""

    def test_replace_simple_text(self, tmp_path):
        """测试简单文本替换。"""
        from agentnexus.utils.docx.ops import replace_text, read_paragraphs

        docx_path = str(tmp_path / "test.docx")
        _create_docx(docx_path, ["Hello World", "测试内容"])

        result = replace_text(docx_path, "Hello", "你好")

        assert result["status"] == "ok"
        assert result["replacements"] >= 1

        # 验证替换结果
        paragraphs = read_paragraphs(docx_path)
        texts = [p["text"] for p in paragraphs]
        assert any("你好" in t for t in texts)

    def test_replace_in_table(self, tmp_path):
        """测试替换表格中的文本。"""
        from agentnexus.utils.docx.ops import replace_text, read_tables

        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_table(docx_path)

        result = replace_text(docx_path, "R0C0", "新内容")
        assert result["status"] == "ok"

    def test_replace_empty_old_text(self, tmp_path):
        """测试空的查找文本。"""
        from agentnexus.utils.docx.ops import replace_text

        docx_path = str(tmp_path / "test.docx")
        _create_docx(docx_path)

        result = replace_text(docx_path, "", "新文本")
        assert result["status"] == "error"

    def test_replace_with_output_path(self, tmp_path):
        """测试替换并输出到新路径。"""
        from agentnexus.utils.docx.ops import replace_text

        docx_path = str(tmp_path / "test.docx")
        out_path = str(tmp_path / "output.docx")
        _create_docx(docx_path, ["原始内容"])

        result = replace_text(docx_path, "原始", "修改后", output_path=out_path)

        assert result["status"] == "ok"
        assert Path(out_path).exists()


# ── 测试：表格编辑 ──

class TestEditTableCell:
    """表格单元格编辑测试。"""

    def test_edit_cell(self, tmp_path):
        """测试编辑表格单元格。"""
        from agentnexus.utils.docx.ops import edit_table_cell, read_tables

        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_table(docx_path)

        result = edit_table_cell(docx_path, 0, 1, 2, "新数据")

        assert result["status"] == "ok"

        # 验证编辑结果
        tables = read_tables(docx_path)
        assert tables[0]["cells"][1][2] == "新数据"

    def test_edit_cell_invalid_table_index(self, tmp_path):
        """测试无效的表格索引。"""
        from agentnexus.utils.docx.ops import edit_table_cell

        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_table(docx_path)

        result = edit_table_cell(docx_path, 99, 0, 0, "数据")
        assert result["status"] == "error"

    def test_edit_cell_invalid_row(self, tmp_path):
        """测试无效的行索引。"""
        from agentnexus.utils.docx.ops import edit_table_cell

        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_table(docx_path)

        result = edit_table_cell(docx_path, 0, 99, 0, "数据")
        assert result["status"] == "error"

    def test_edit_cell_overflow_warning(self, tmp_path):
        """测试内容溢出时的警告。"""
        from agentnexus.utils.docx.ops import edit_table_cell

        docx_path = str(tmp_path / "test.docx")
        _create_docx_with_table(docx_path)

        # 写入很长的文本
        long_text = "这是一段非常长的文本" * 20
        result = edit_table_cell(docx_path, 0, 0, 0, long_text)

        assert result["status"] == "ok"
        # 可能有溢出警告
        if "warnings" in result:
            assert any("超出" in w for w in result["warnings"])


# ── 测试：插入操作 ──

class TestInsertOperations:
    """插入操作测试。"""

    def test_insert_paragraph(self, tmp_path):
        """测试插入段落。"""
        from agentnexus.utils.docx.ops import insert_paragraph, read_paragraphs

        docx_path = str(tmp_path / "test.docx")
        _create_docx(docx_path, ["段落1", "段落2"])

        result = insert_paragraph(docx_path, 0, "新插入的段落")

        assert result["status"] == "ok"

        paragraphs = read_paragraphs(docx_path)
        texts = [p["text"] for p in paragraphs]
        assert "新插入的段落" in texts

    def test_insert_paragraph_at_beginning(self, tmp_path):
        """测试在文档开头插入段落。"""
        from agentnexus.utils.docx.ops import insert_paragraph, read_paragraphs

        docx_path = str(tmp_path / "test.docx")
        _create_docx(docx_path, ["原始段落"])

        result = insert_paragraph(docx_path, -1, "开头段落")

        assert result["status"] == "ok"

    def test_insert_paragraph_invalid_index(self, tmp_path):
        """测试无效的段落索引。"""
        from agentnexus.utils.docx.ops import insert_paragraph

        docx_path = str(tmp_path / "test.docx")
        _create_docx(docx_path)

        result = insert_paragraph(docx_path, 999, "内容")
        assert result["status"] == "error"

    def test_insert_table(self, tmp_path):
        """测试插入表格。"""
        from agentnexus.utils.docx.ops import insert_table, read_tables

        docx_path = str(tmp_path / "test.docx")
        _create_docx(docx_path, ["段落1", "段落2"])

        data = [["A", "B", "C"], ["1", "2", "3"]]
        result = insert_table(docx_path, 0, data)

        assert result["status"] == "ok"
        assert result["rows"] == 2
        assert result["cols"] == 3

    def test_insert_table_empty_data(self, tmp_path):
        """测试空表格数据。"""
        from agentnexus.utils.docx.ops import insert_table

        docx_path = str(tmp_path / "test.docx")
        _create_docx(docx_path)

        result = insert_table(docx_path, 0, [])
        assert result["status"] == "error"


# ── 测试：页面设置 ──

class TestPageSettings:
    """页面设置测试。"""

    def test_set_margins(self, tmp_path):
        """测试设置页面边距。"""
        from agentnexus.utils.docx.ops import set_page_margins
        from agentnexus.utils.docx.analyzer import analyze

        docx_path = str(tmp_path / "test.docx")
        _create_docx(docx_path)

        result = set_page_margins(docx_path, top_mm=20, left_mm=25)

        assert result["status"] == "ok"

        # 验证边距已修改
        constraints = analyze(docx_path)
        page = constraints.default_page
        assert 19 < page.margin_top_mm < 21
        assert 24 < page.margin_left_mm < 26


# ── 测试：文件操作 ──

class TestFileOperations:
    """文件操作测试。"""

    def test_save_as(self, tmp_path):
        """测试另存为。"""
        from agentnexus.utils.docx.ops import save_as

        src = str(tmp_path / "source.docx")
        dst = str(tmp_path / "output.docx")
        _create_docx(src)

        result = save_as(src, dst)

        assert result["status"] == "ok"
        assert Path(dst).exists()

    def test_save_as_nonexistent(self, tmp_path):
        """测试另存为不存在的文件。"""
        from agentnexus.utils.docx.ops import save_as

        result = save_as("/nonexistent.docx", str(tmp_path / "out.docx"))
        assert result["status"] == "error"


# ── 测试：验证操作 ──

class TestValidation:
    """格式验证测试。"""

    def test_validate_simple_doc(self, tmp_path):
        """测试验证简单文档。"""
        from agentnexus.utils.docx.ops import validate

        docx_path = str(tmp_path / "test.docx")
        _create_docx(docx_path)

        issues = validate(docx_path)
        # 简单文档应该没有严重问题
        errors = [i for i in issues if i["severity"] == "error"]
        assert len(errors) == 0

    def test_validate_nonexistent(self):
        """测试验证不存在的文件。"""
        from agentnexus.utils.docx.ops import validate

        issues = validate("/nonexistent.docx")
        assert len(issues) > 0
        assert issues[0]["severity"] == "error"
